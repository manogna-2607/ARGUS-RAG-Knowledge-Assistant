#!/usr/bin/env python3
"""ARGUS document extraction worker.

This process accepts a local source file path and document type, then returns
structured JSON to the Node.js ingestion layer. It never prints diagnostics to
stdout so the caller can safely parse the response.
"""

from __future__ import annotations

import argparse
import io
import json
import sys
from pathlib import Path
from typing import Any


class ExtractionError(Exception):
    """A safe, user-facing extraction error."""


class ScannedPdfError(ExtractionError):
    """Raised when every page lacks machine-readable text."""


def non_empty_pages(pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Remove pages that contain no extractable text, preserving page numbers."""
    return [
        {"text": str(page["text"]).strip(), "pageNumber": page["pageNumber"]}
        for page in pages
        if str(page.get("text", "")).strip()
    ]


def extract_text_from_pdf(file_bytes: bytes) -> dict[str, Any]:
    """Extract page-bounded PDF text with PyMuPDF and safe fallback parsers.

    PyMuPDF (fitz/pymupdf) is the primary parser because it is fast, tolerant
    of unusual layouts, and preserves source page boundaries. pdfplumber and
    pypdf are only used if PyMuPDF is unavailable or rejects a malformed PDF.
    """
    parser_errors: list[str] = []

    # Primary parser: PyMuPDF
    try:
        try:
            import pymupdf as fitz  # PyMuPDF >= 1.24 modern import
        except ImportError:
            import fitz  # pragma: no cover - compatibility with old PyMuPDF

        document = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            if document.is_encrypted and not document.authenticate(""):
                raise ExtractionError("This PDF is encrypted and requires a password.")

            pages = []
            for page_index in range(document.page_count):
                page = document.load_page(page_index)
                # sort=True provides reading order where available and is more
                # reliable for common multicolumn PDFs than raw extraction.
                text = page.get_text("text", sort=True)
                pages.append({"text": text, "pageNumber": page_index + 1})

            pages = non_empty_pages(pages)
            if not pages:
                raise ScannedPdfError(
                    "This PDF appears to contain scanned/image-only pages. OCR is required for these pages."
                )

            return {"pages": pages, "parser": "PyMuPDF", "pageCount": document.page_count}
        finally:
            document.close()
    except ScannedPdfError:
        # An image-only PDF will not become text-searchable via fallback parsers.
        raise
    except ExtractionError:
        raise
    except Exception as error:  # deliberately continue to fallback parsers
        parser_errors.append(f"PyMuPDF: {error}")

    # Fallback parser: pdfplumber
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as document:
            total_pages = len(document.pages)
            pages = []
            for page_index, page in enumerate(document.pages):
                # layout=True retains useful visual line grouping for headings.
                text = page.extract_text(layout=True) or page.extract_text() or ""
                pages.append({"text": text, "pageNumber": page_index + 1})

        pages = non_empty_pages(pages)
        if not pages:
            raise ScannedPdfError(
                "This PDF appears to contain scanned/image-only pages. OCR is required for these pages."
            )

        return {"pages": pages, "parser": "pdfplumber", "pageCount": total_pages}
    except ScannedPdfError:
        raise
    except Exception as error:
        parser_errors.append(f"pdfplumber: {error}")

    # Final fallback: pypdf
    try:
        from pypdf import PdfReader

        document = PdfReader(io.BytesIO(file_bytes), strict=False)
        if document.is_encrypted and not document.decrypt(""):
            raise ExtractionError("This PDF is encrypted and requires a password.")

        pages = []
        for page_index, page in enumerate(document.pages):
            pages.append({"text": page.extract_text() or "", "pageNumber": page_index + 1})

        pages = non_empty_pages(pages)
        if not pages:
            raise ScannedPdfError(
                "This PDF appears to contain scanned/image-only pages. OCR is required for these pages."
            )

        return {"pages": pages, "parser": "pypdf", "pageCount": len(document.pages)}
    except ScannedPdfError:
        raise
    except ExtractionError:
        raise
    except Exception as error:
        parser_errors.append(f"pypdf: {error}")

    error_summary = "; ".join(parser_errors)
    raise ExtractionError(f"Unable to extract this PDF. Parser details: {error_summary}")


def extract_text_from_docx(file_bytes: bytes) -> dict[str, Any]:
    """Extract DOCX paragraphs and accessible table cells with python-docx."""
    try:
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
        sections: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            # Preserve hierarchy for heading styles so chunking respects sections.
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                sections.append(f"\n{text}\n")
            else:
                sections.append(text)

        for table_number, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                sections.append(f"\nTable {table_number}:\n" + "\n".join(rows))

        text = "\n\n".join(sections).strip()
        if not text:
            raise ExtractionError("The DOCX file contains no readable paragraphs or table text.")

        return {
            "pages": [{"text": text, "pageNumber": None}],
            "parser": "python-docx",
            "pageCount": 0,
        }
    except ExtractionError:
        raise
    except Exception as error:
        raise ExtractionError(f"Unable to parse DOCX content: {error}") from error


def extract_text_from_txt(file_bytes: bytes) -> dict[str, Any]:
    """Decode text safely, preferring UTF-8 while accommodating common files."""
    text = ""
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError as error:
            last_error = error
    else:
        raise ExtractionError(f"Unable to decode TXT file: {last_error}")

    if not text.strip():
        raise ExtractionError("The TXT file contains no readable text.")

    return {
        "pages": [{"text": text, "pageNumber": None}],
        "parser": "text-decoder",
        "pageCount": 0,
    }


def extract_document(file_bytes: bytes, document_type: str) -> dict[str, Any]:
    """Select an extractor, then return non-empty page-bounded text."""
    if document_type == "pdf":
        result = extract_text_from_pdf(file_bytes)
    elif document_type == "docx":
        result = extract_text_from_docx(file_bytes)
    elif document_type == "txt":
        result = extract_text_from_txt(file_bytes)
    else:
        raise ExtractionError(f"Unsupported document type: {document_type}")

    pages = non_empty_pages(result["pages"])
    text = "\n\n".join(page["text"] for page in pages).strip()
    if not text:
        raise ExtractionError("No readable text content could be extracted from this document.")

    return {
        "ok": True,
        "pages": pages,
        "text": text,
        "pageCount": result["pageCount"],
        "extractedCharacters": len(text),
        "parser": result["parser"],
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="ARGUS document extraction worker")
    parser.add_argument("--input", required=True, help="Path to uploaded document bytes")
    parser.add_argument("--type", required=True, choices=("pdf", "docx", "txt"))
    args = parser.parse_args()

    try:
        file_bytes = Path(args.input).read_bytes()
        result = extract_document(file_bytes, args.type)
        print(json.dumps(result, ensure_ascii=False))
    except ScannedPdfError as error:
        print(json.dumps({"ok": False, "code": "SCANNED_PDF", "error": str(error)}))
        sys.exit(2)
    except ExtractionError as error:
        print(json.dumps({"ok": False, "code": "EXTRACTION_FAILED", "error": str(error)}))
        sys.exit(2)
    except Exception as error:  # final protection against runtime worker failure
        print(json.dumps({"ok": False, "code": "EXTRACTION_FAILED", "error": f"Extraction worker failed: {error}"}))
        sys.exit(3)


if __name__ == "__main__":
    main()
